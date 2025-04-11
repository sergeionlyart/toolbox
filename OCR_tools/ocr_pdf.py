#!/usr/bin/env python3
# python3 ocr_pdf.py
""" 
⸻

📄 OCR PDF Processor

OCR PDF Processor — это Python-скрипт, который позволяет автоматически распознавать текст и изображения в PDF-документах с помощью OCR-сервиса от Mistral AI, и сохранять результаты в виде:
	•	📑 DOCX-файла с интегрированными изображениями и распознанным текстом.
	•	🗄 JSON-файла с детализированными результатами OCR (текст и изображения в base64).

⚙️ Основные возможности:
	•	Поддерживает обработку как локальных, так и удалённых PDF-файлов.
	•	Автоматически загружает и обрабатывает PDF-документы через API Mistral OCR.
	•	Конвертирует результаты OCR в удобный для работы DOCX-формат.
	•	Сохраняет подробный JSON-отчёт с результатами OCR.

🚀 Как использовать:

Запуск скрипта осуществляется командой:

python ocr_pdf.py "<путь_или_URL_к_PDF>"

Пример использования для локального файла:

python ocr_pdf.py "file:///путь/к/документу.pdf"

Результаты сохраняются в директориях:
	•	DOCX-файл: docs/processed/
	•	JSON-файл: docs/JSON_data/

🛠 Зависимости:
	•	mistralai
	•	requests
	•	python-docx
	•	Pillow

Перед запуском установите их командой:

pip install mistralai requests python-docx Pillow

Также убедитесь, что установлен ваш API-ключ Mistral AI.

⸻

Этот инструмент полезен для быстрого преобразования PDF-документов в редактируемые форматы для последующей работы с текстом и изображениями.



"""

#!/usr/bin/env python3

import sys
import logging
import requests
import shutil
from urllib.parse import urlparse, unquote
from pathlib import Path
from mistralai import Mistral, DocumentURLChunk, OCRResponse
import re
import base64
import json
from io import BytesIO
from docx import Document
from PIL import Image  # (по желанию используйте для проверки)
from docx.image.exceptions import UnrecognizedImageError
from multi_tenant_config import MISTRAL_API_KEY

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')

# ---------------------------
# Вспомогательные функции
# ---------------------------

def replace_images_in_markdown(markdown_str: str, images_dict: dict) -> str:
    """
    Заменяет плейсхолдеры вида ![imgId](imgId) на base64-encoded изображения.
    """
    for img_name, base64_str in images_dict.items():
        markdown_str = markdown_str.replace(
            f"![{img_name}]({img_name})",
            f"![{img_name}](data:image/png;base64,{base64_str})"
        )
    return markdown_str

def get_combined_markdown(ocr_response: OCRResponse) -> str:
    """
    Объединяет текст и картинки каждой страницы OCR-ответа в единый Markdown.
    """
    markdowns = []
    for page in ocr_response.pages:
        image_data = {img.id: img.image_base64 for img in page.images}
        markdowns.append(replace_images_in_markdown(page.markdown, image_data))
    return "\n\n".join(markdowns)

def save_markdown_with_images_to_docx(markdown_text: str, filename: str):
    """
    Интегрированный подход, позволяющий на основе markdown-текста с base64-картинками
    формировать DOCX-документ с интегрированными изображениями.

    Аргументы:
      markdown_text: итоговый markdown-текст (с base64 вместо ссылок на картинки)
      filename: путь и имя файла (например 'output.docx')
    """
    import re
    import base64
    from io import BytesIO
    from docx import Document

    # Ищем все картинки вида ![alt](data:image/png;base64,....)
    image_pattern = r'!\[(.*?)\]\(\s*(data:image/(?:png|jpg|jpeg|gif);base64,([^)]*))\)'

    document = Document()

    # Разделяем текст на строки, чтобы создавать параграфы, и обрабатываем каждую строку
    lines = markdown_text.split('\n')
    for line in lines:
        # Ищем все картинки в строке
        matches = re.findall(image_pattern, line)

        if matches:
            # Если строка содержит изображение, то обрабатываем
            text_only = re.sub(image_pattern, '', line).strip()
            if text_only:
                document.add_paragraph(text_only)

            # Для каждой картинки создаём абзац + run
            from PIL import Image

            for alt_text, full_data_string, b64_data in matches:
                # Убираем префикс data:image/...;base64, если он есть
                if full_data_string.startswith("data:image"):
                    b64_data_clean = full_data_string.split("base64,")[-1]
                else:
                    b64_data_clean = b64_data

                image_data = base64.b64decode(b64_data_clean)
                image_stream = BytesIO(image_data)

                try:
                    with Image.open(image_stream) as img:
                        image_format = img.format if img.format else 'PNG'
                        image_stream.seek(0)

                        run = document.add_paragraph().add_run()
                        run.add_picture(image_stream)
                except Exception as img_err:
                    logging.error(f"Ошибка обработки изображения: {img_err}")
        else:
            # Если в строке нет изображений, добавляем её как обычный текст
            document.add_paragraph(line)

    # Сохраняем документ в файл
    document.save(filename)
    logging.info(f"DOCX файл успешно сохранён как '{filename}'.")

# ---------------------------
# Основная логика с поддержкой локальных и сетевых файлов
# ---------------------------

def process_pdf_from_url(pdf_url: str, output_dir: Path):
    """
    1. Получаем PDF (локальный/удалённый).
    2. Загружаем на Mistral.
    3. OCR -> Получаем текст + картинки
    4. Сохраняем JSON в dev_lab/docs/JSON_data/
    5. Формируем Markdown + DOCX
    """
    try:
        logging.info(f"Starting processing for URL: {pdf_url}")
        client = Mistral(api_key=MISTRAL_API_KEY)

        parsed_url = urlparse(pdf_url)
        pdf_name = Path(unquote(parsed_url.path)).stem
        output_filename = output_dir / f"{pdf_name}_ocr_processed.docx"

        local_pdf_path = Path("temp.pdf")

        # 1. Обработка локального или удаленного файла
        if parsed_url.scheme == 'file':
            local_file_path = Path(unquote(parsed_url.path))
            logging.info(f"Copying local PDF from {local_file_path}...")
            shutil.copy(local_file_path, local_pdf_path)
            logging.info("PDF copied successfully.")
        else:
            logging.info(f"Downloading PDF from {pdf_url}...")
            resp = requests.get(pdf_url, timeout=30)
            resp.raise_for_status()
            local_pdf_path.write_bytes(resp.content)
            logging.info("PDF downloaded successfully.")

        # 2. Загрузка файла на Mistral
        logging.info("Uploading PDF to Mistral...")
        uploaded_file = client.files.upload(
            file={
                "file_name": local_pdf_path.stem,
                "content": local_pdf_path.read_bytes(),
            },
            purpose="ocr",
        )
        logging.info("PDF uploaded successfully.")

        # 3. Получение подписанного URL для OCR
        logging.info("Getting signed URL...")
        signed_url = client.files.get_signed_url(file_id=uploaded_file.id, expiry=1)
        logging.info(f"Signed URL obtained: {signed_url.url}")

        # 4. Выполнение OCR
        logging.info("Starting OCR processing...")
        pdf_response = client.ocr.process(
            document=DocumentURLChunk(document_url=signed_url.url),
            model="mistral-ocr-latest",
            include_image_base64=True
        )
        logging.info("OCR processing completed.")

        # Дополнительно сохраняем JSON-результат
        json_data_dir = Path(__file__).parent / "docs/JSON_data"
        json_data_dir.mkdir(parents=True, exist_ok=True)
        json_path = json_data_dir / f"{pdf_name}.json"

        response_dict = pdf_response.model_dump()
        with json_path.open('w', encoding='utf-8') as f:
            json.dump(response_dict, f, ensure_ascii=False, indent=2)
        logging.info(f"JSON data saved to '{json_path}'")

        # 5. Генерация markdown и сохранение DOCX
        logging.info("Generating markdown from OCR response...")
        combined_md = get_combined_markdown(pdf_response)
        save_markdown_with_images_to_docx(str(combined_md), str(output_filename))

    except Exception as e:
        logging.exception(f"An error occurred: {e}")
    finally:
        if local_pdf_path.exists():
            local_pdf_path.unlink(missing_ok=True)
            logging.info("Temporary PDF file removed.")

    logging.info(f"Process finished. Output at '{output_filename}'.")

# ---------------------------
# Входная точка скрипта
# ---------------------------

if __name__ == "__main__":
    if len(sys.argv) != 2:
        logging.error("Usage: python ocr_pdf.py <absolute_pdf_url>")
        sys.exit(1)

    pdf_url = sys.argv[1]
    output_directory = Path(__file__).parent / "docs/processed"
    process_pdf_from_url(pdf_url, output_directory)